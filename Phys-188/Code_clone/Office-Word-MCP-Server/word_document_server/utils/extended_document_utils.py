"""
Extended document utilities for Word Document Server.
"""
from typing import Dict, List, Any, Tuple
from docx import Document


def get_paragraph_text(doc_path: str, paragraph_index: int) -> Dict[str, Any]:
    """
    Get text from a specific paragraph in a Word document.
    
    Args:
        doc_path: Path to the Word document
        paragraph_index: Index of the paragraph to extract (0-based)
    
    Returns:
        Dictionary with paragraph text and metadata
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        
        # Check if paragraph index is valid
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {"error": f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."}
        
        paragraph = doc.paragraphs[paragraph_index]
        
        return {
            "index": paragraph_index,
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "is_heading": paragraph.style.name.startswith("Heading") if paragraph.style else False
        }
    except Exception as e:
        return {"error": f"Failed to get paragraph text: {str(e)}"}


def find_text(doc_path: str, text_to_find: str, match_case: bool = True, whole_word: bool = False) -> Dict[str, Any]:
    """
    Find all occurrences of specific text in a Word document.
    
    Args:
        doc_path: Path to the Word document
        text_to_find: Text to search for
        match_case: Whether to perform case-sensitive search
        whole_word: Whether to match whole words only
    
    Returns:
        Dictionary with search results
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    if not text_to_find:
        return {"error": "Search text cannot be empty"}
    
    try:
        doc = Document(doc_path)
        results = {
            "query": text_to_find,
            "match_case": match_case,
            "whole_word": whole_word,
            "occurrences": [],
            "total_count": 0
        }
        
        # Search in paragraphs
        for i, para in enumerate(doc.paragraphs):
            # Prepare text for comparison
            para_text = para.text
            search_text = text_to_find
            
            if not match_case:
                para_text = para_text.lower()
                search_text = search_text.lower()
            
            # Find all occurrences (simple implementation)
            start_pos = 0
            while True:
                if whole_word:
                    # For whole word search, we need to check word boundaries
                    words = para_text.split()
                    found = False
                    for word_idx, word in enumerate(words):
                        if (word == search_text or 
                            (not match_case and word.lower() == search_text.lower())):
                            results["occurrences"].append({
                                "paragraph_index": i,
                                "position": word_idx,
                                "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                            })
                            results["total_count"] += 1
                            found = True
                    
                    # Break after checking all words
                    break
                else:
                    # For substring search
                    pos = para_text.find(search_text, start_pos)
                    if pos == -1:
                        break
                    
                    results["occurrences"].append({
                        "paragraph_index": i,
                        "position": pos,
                        "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                    })
                    results["total_count"] += 1
                    start_pos = pos + len(search_text)
        
        # Search in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        # Prepare text for comparison
                        para_text = para.text
                        search_text = text_to_find
                        
                        if not match_case:
                            para_text = para_text.lower()
                            search_text = search_text.lower()
                        
                        # Find all occurrences (simple implementation)
                        start_pos = 0
                        while True:
                            if whole_word:
                                # For whole word search, check word boundaries
                                words = para_text.split()
                                found = False
                                for word_idx, word in enumerate(words):
                                    if (word == search_text or 
                                        (not match_case and word.lower() == search_text.lower())):
                                        results["occurrences"].append({
                                            "location": f"Table {table_idx}, Row {row_idx}, Column {col_idx}",
                                            "position": word_idx,
                                            "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                                        })
                                        results["total_count"] += 1
                                        found = True
                                
                                # Break after checking all words
                                break
                            else:
                                # For substring search
                                pos = para_text.find(search_text, start_pos)
                                if pos == -1:
                                    break
                                
                                results["occurrences"].append({
                                    "location": f"Table {table_idx}, Row {row_idx}, Column {col_idx}",
                                    "position": pos,
                                    "context": para.text[:100] + ("..." if len(para.text) > 100 else "")
                                })
                                results["total_count"] += 1
                                start_pos = pos + len(search_text)
        
        return results
    except Exception as e:
        return {"error": f"Failed to search for text: {str(e)}"}
