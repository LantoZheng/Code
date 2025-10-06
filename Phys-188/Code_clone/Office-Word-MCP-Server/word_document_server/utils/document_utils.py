"""
Document utility functions for Word Document Server.
"""
import json
from typing import Dict, List, Any
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs


def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document, skipping Table of Contents (TOC) paragraphs.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Search in paragraphs
    for para in doc.paragraphs:
        # Skip TOC paragraphs
        if para.style and para.style.name.startswith("TOC"):
            continue
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Skip TOC paragraphs in tables
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    
    return count


def get_document_xml(doc_path: str) -> str:
    """Extract and return the raw XML structure of the Word document (word/document.xml)."""
    import os
    import zipfile
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        with zipfile.ZipFile(doc_path) as docx_zip:
            with docx_zip.open('word/document.xml') as xml_file:
                return xml_file.read().decode('utf-8')
    except Exception as e:
        return f"Failed to extract XML: {str(e)}"


def insert_header_near_text(doc_path: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search."""
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        new_para = doc.add_paragraph(header_title, style=header_style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert header: {str(e)}"


def insert_line_or_paragraph_near_text(doc_path: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """
    Insert a new line or paragraph (with specified or matched style) before or after the target paragraph.
    You can specify the target by text (first match) or by paragraph index.
    Skips paragraphs whose style name starts with 'TOC' if using text search.
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine style: use provided or match target
        style = line_style if line_style else para.style
        new_para = doc.add_paragraph(line_text, style=style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Line/paragraph inserted {position} paragraph (index {anchor_index}) with style '{style}'."
        else:
            return f"Line/paragraph inserted {position} the target paragraph with style '{style}'."
    except Exception as e:
        return f"Failed to insert line/paragraph: {str(e)}"


def insert_numbered_list_near_text(doc_path: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None) -> str:
    """
    Insert a numbered list before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search.
    Args:
        doc_path: Path to the Word document
        target_text: Text to search for in paragraphs (optional if using index)
        list_items: List of strings, each as a list item
        position: 'before' or 'after' (default: 'after')
        target_paragraph_index: Optional paragraph index to use as anchor
    Returns:
        Status message
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Robust style selection for numbered list
        style_name = None
        for candidate in ['List Number', 'List Paragraph', 'Normal']:
            try:
                _ = doc.styles[candidate]
                style_name = candidate
                break
            except KeyError:
                continue
        if not style_name:
            style_name = None  # fallback to default
        new_paras = []
        for item in (list_items or []):
            p = doc.add_paragraph(item, style=style_name)
            new_paras.append(p)
        # Move the new paragraphs to the correct position
        for p in reversed(new_paras):
            if position == 'before':
                para._element.addprevious(p._element)
            else:
                para._element.addnext(p._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Numbered list inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Numbered list inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert numbered list: {str(e)}"


def is_toc_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de tabla de contenido (TOC)."""
    return para.style and para.style.name.upper().startswith("TOC")


def is_heading_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de encabezado (Heading 1, Heading 2, etc)."""
    return para.style and para.style.name.lower().startswith("heading")


# --- Helper: Get style name from a <w:p> element ---
def get_paragraph_style(el):
    from docx.oxml.ns import qn
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and 'w:val' in pStyle.attrib:
            return pStyle.attrib['w:val']
    return None

# --- Main: Delete everything under a header until next heading/TOC ---
def delete_block_under_header(doc, header_text):
    """
    Remove all elements (paragraphs, tables, etc.) after the header (by text) and before the next heading/TOC (by style).
    Returns: (header_element, elements_removed)
    """
    # Find the header paragraph by text (like delete_paragraph finds by index)
    header_para = None
    header_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == header_text.strip().lower():
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return None, 0
    
    # Find the next heading/TOC paragraph to determine the end of the block
    end_idx = None
    for i in range(header_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style and para.style.name.lower().startswith(('heading', 'título', 'toc')):
            end_idx = i
            break
    
    # If no next heading found, delete until end of document
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    # Remove paragraphs by index (like delete_paragraph does)
    removed_count = 0
    for i in range(header_idx + 1, end_idx):
        if i < len(doc.paragraphs):  # Safety check
            para = doc.paragraphs[header_idx + 1]  # Always remove the first paragraph after header
            p = para._p
            p.getparent().remove(p)
            removed_count += 1
    
    return header_para._p, removed_count

# --- Usage in replace_paragraph_block_below_header ---
def replace_paragraph_block_below_header(
    doc_path: str,
    header_text: str,
    new_paragraphs: list,
    detect_block_end_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Reemplaza todo el contenido debajo de una cabecera (por texto), hasta el siguiente encabezado/TOC (por estilo).
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    
    doc = Document(doc_path)
    
    # Find the header paragraph first
    header_para = None
    header_idx = None
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip().lower()
        is_toc = is_toc_paragraph(para)
        if para_text == header_text.strip().lower() and not is_toc:
            header_para = para
            header_idx = i
            break
    
    if header_para is None:
        return f"Header '{header_text}' not found in document."
    
    # Delete everything under the header using the same document instance
    header_el, removed_count = delete_block_under_header(doc, header_text)
    
    # Now insert new paragraphs after the header (which should still be in the document)
    style_to_use = new_paragraph_style or "Normal"
    
    # Find the header again after deletion (it should still be there)
    current_para = header_para
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        current_para._element.addnext(new_para._element)
        current_para = new_para
    
    doc.save(doc_path)
    return f"Replaced content under '{header_text}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {removed_count} elements."


def replace_block_between_manual_anchors(
    doc_path: str,
    start_anchor_text: str,
    new_paragraphs: list,
    end_anchor_text: str = None,
    match_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all content (paragraphs, tables, etc.) between start_anchor_text and end_anchor_text (or next logical header if not provided).
    If end_anchor_text is None, deletes until next visually distinct paragraph (bold, all caps, or different font size), or end of document.
    Inserts new_paragraphs after the start anchor.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    doc = Document(doc_path)
    body = doc.element.body
    elements = list(body)
    start_idx = None
    end_idx = None
    # Find start anchor
    for i, el in enumerate(elements):
        if el.tag == CT_P.tag:
            p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
            if match_fn:
                if match_fn(p_text, el):
                    start_idx = i
                    break
            elif p_text == start_anchor_text.strip():
                start_idx = i
                break
    if start_idx is None:
        return f"Start anchor '{start_anchor_text}' not found."
    # Find end anchor
    if end_anchor_text:
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                if match_fn:
                    if match_fn(p_text, el, is_end=True):
                        end_idx = i
                        break
                elif p_text == end_anchor_text.strip():
                    end_idx = i
                    break
    else:
        # Heuristic: next visually distinct paragraph (bold, all caps, or different font size), or end of document
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == CT_P.tag:
                # Check for bold, all caps, or font size
                runs = [node for node in el.iter() if node.tag.endswith('}r')]
                for run in runs:
                    rpr = run.find(qn('w:rPr'))
                    if rpr is not None:
                        if rpr.find(qn('w:b')) is not None or rpr.find(qn('w:caps')) is not None or rpr.find(qn('w:sz')) is not None:
                            end_idx = i
                            break
                if end_idx is not None:
                    break
    # Mark elements for removal
    to_remove = []
    for i in range(start_idx + 1, end_idx if end_idx is not None else len(elements)):
        to_remove.append(elements[i])
    for el in to_remove:
        body.remove(el)
    doc.save(doc_path)
    # Reload and find start anchor for insertion
    doc = Document(doc_path)
    paras = doc.paragraphs
    anchor_idx = None
    for i, para in enumerate(paras):
        if para.text.strip() == start_anchor_text.strip():
            anchor_idx = i
            break
    if anchor_idx is None:
        return f"Start anchor '{start_anchor_text}' not found after deletion (unexpected)."
    anchor_para = paras[anchor_idx]
    style_to_use = new_paragraph_style or "Normal"
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        anchor_para._element.addnext(new_para._element)
        anchor_para = new_para
    doc.save(doc_path)
    return f"Replaced content between '{start_anchor_text}' and '{end_anchor_text or 'next logical header'}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {len(to_remove)} elements."
