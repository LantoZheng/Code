"""
Table-related operations for Word Document Server.
"""
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    
    Args:
        cell: The cell to modify
        **kwargs: Border properties (top, bottom, left, right, val, color)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for key, value in kwargs.items():
        if key in ['top', 'left', 'bottom', 'right']:
            tag = 'w:{}'.format(key)
            
            element = OxmlElement(tag)
            element.set(qn('w:val'), kwargs.get('val', 'single'))
            element.set(qn('w:sz'), kwargs.get('sz', '4'))
            element.set(qn('w:space'), kwargs.get('space', '0'))
            element.set(qn('w:color'), kwargs.get('color', 'auto'))
            
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
                
            tcBorders.append(element)


def apply_table_style(table, has_header_row=False, border_style=None, shading=None):
    """
    Apply formatting to a table.
    
    Args:
        table: The table to format
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick')
        shading: 2D list of cell background colors (by row and column)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Format header row if requested
        if has_header_row and table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Apply border style if specified
        if border_style:
            val_map = {
                'none': 'nil',
                'single': 'single',
                'double': 'double',
                'thick': 'thick'
            }
            val = val_map.get(border_style.lower(), 'single')
            
            # Apply to all cells
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top=True,
                        bottom=True,
                        left=True,
                        right=True,
                        val=val,
                        color="000000"
                    )
        
        # Apply cell shading if specified
        if shading:
            for i, row_colors in enumerate(shading):
                if i >= len(table.rows):
                    break
                for j, color in enumerate(row_colors):
                    if j >= len(table.rows[i].cells):
                        break
                    try:
                        # Apply shading to cell
                        cell = table.rows[i].cells[j]
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        # Skip if color format is invalid
                        pass
        
        return True
    except Exception:
        return False


def copy_table(source_table, target_doc):
    """
    Copy a table from one document to another.
    
    Args:
        source_table: The table to copy
        target_doc: The document to copy the table to
        
    Returns:
        The new table in the target document
    """
    # Create a new table with the same dimensions
    new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    
    # Try to apply the same style
    try:
        if source_table.style:
            new_table.style = source_table.style
    except:
        # Fall back to default grid style
        try:
            new_table.style = 'Table Grid'
        except:
            pass
    
    # Copy cell contents
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if paragraph.text:
                    new_table.cell(i, j).text = paragraph.text
    
    return new_table


def set_cell_shading(cell, fill_color=None, pattern="clear", pattern_color="auto"):
    """
    Apply shading/filling to a table cell.
    
    Args:
        cell: The table cell to format
        fill_color: Background color (hex string like "FF0000" or RGBColor)
        pattern: Shading pattern ("clear", "solid", "pct10", "pct20", etc.)
        pattern_color: Pattern color for patterned fills
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Get or create table cell properties
        tc_pr = cell._tc.get_or_add_tcPr()
        
        # Remove existing shading
        existing_shd = tc_pr.find(qn('w:shd'))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)
        
        # Create shading element
        shd_attrs = {
            'w:val': pattern,
            'w:color': pattern_color if pattern_color != "auto" else "auto"
        }
        
        # Set fill color
        if fill_color:
            if isinstance(fill_color, str):
                # Hex color string - remove # if present
                fill_color = fill_color.lstrip('#').upper()
                if len(fill_color) == 6:  # Valid hex color
                    shd_attrs['w:fill'] = fill_color
            elif isinstance(fill_color, RGBColor):
                # RGBColor object
                hex_color = f"{fill_color.r:02X}{fill_color.g:02X}{fill_color.b:02X}"
                shd_attrs['w:fill'] = hex_color
        
        # Build XML string
        attr_str = ' '.join([f'{k}="{v}"' for k, v in shd_attrs.items()])
        shd_xml = f'<w:shd {nsdecls("w")} {attr_str}/>'
        
        # Parse and append shading element
        shading_elm = parse_xml(shd_xml)
        tc_pr.append(shading_elm)
        
        return True
        
    except Exception as e:
        print(f"Error setting cell shading: {e}")
        return False


def apply_alternating_row_shading(table, color1="FFFFFF", color2="F2F2F2"):
    """
    Apply alternating row colors for better readability.
    
    Args:
        table: The table to format
        color1: Color for odd rows (hex string)
        color2: Color for even rows (hex string)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        for i, row in enumerate(table.rows):
            fill_color = color1 if i % 2 == 0 else color2
            for cell in row.cells:
                set_cell_shading(cell, fill_color=fill_color)
        return True
    except Exception as e:
        print(f"Error applying alternating row shading: {e}")
        return False


def highlight_header_row(table, header_color="4472C4", text_color="FFFFFF"):
    """
    Apply special shading to header row.
    
    Args:
        table: The table to format
        header_color: Background color for header (hex string)
        text_color: Text color for header (hex string)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if table.rows:
            for cell in table.rows[0].cells:
                # Apply background shading
                set_cell_shading(cell, fill_color=header_color)
                
                # Apply text formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        if text_color and text_color != "auto":
                            # Convert hex to RGB
                            try:
                                text_color = text_color.lstrip('#')
                                r = int(text_color[0:2], 16)
                                g = int(text_color[2:4], 16)
                                b = int(text_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                            except:
                                pass  # Skip if color format is invalid
        return True
    except Exception as e:
        print(f"Error highlighting header row: {e}")
        return False


def set_cell_shading_by_position(table, row_index, col_index, fill_color, pattern="clear"):
    """
    Apply shading to a specific cell by row/column position.
    
    Args:
        table: The table containing the cell
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        fill_color: Background color (hex string)
        pattern: Shading pattern
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if (0 <= row_index < len(table.rows) and 
            0 <= col_index < len(table.rows[row_index].cells)):
            cell = table.rows[row_index].cells[col_index]
            return set_cell_shading(cell, fill_color=fill_color, pattern=pattern)
        else:
            return False
    except Exception as e:
        print(f"Error setting cell shading by position: {e}")
        return False


def merge_cells(table, start_row, start_col, end_row, end_col):
    """
    Merge cells in a rectangular area.
    
    Args:
        table: The table containing cells to merge
        start_row: Starting row index (0-based)
        start_col: Starting column index (0-based)
        end_row: Ending row index (0-based, inclusive)
        end_col: Ending column index (0-based, inclusive)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Validate indices
        if (start_row < 0 or start_col < 0 or end_row < 0 or end_col < 0 or
            start_row >= len(table.rows) or end_row >= len(table.rows) or
            start_row > end_row or start_col > end_col):
            return False
        
        # Check if all rows have enough columns
        for row_idx in range(start_row, end_row + 1):
            if (start_col >= len(table.rows[row_idx].cells) or 
                end_col >= len(table.rows[row_idx].cells)):
                return False
        
        # Get the start and end cells
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        
        # Merge the cells
        start_cell.merge(end_cell)
        
        return True
        
    except Exception as e:
        print(f"Error merging cells: {e}")
        return False


def merge_cells_horizontal(table, row_index, start_col, end_col):
    """
    Merge cells horizontally in a single row.
    
    Args:
        table: The table containing cells to merge
        row_index: Row index (0-based)
        start_col: Starting column index (0-based)
        end_col: Ending column index (0-based, inclusive)
        
    Returns:
        True if successful, False otherwise
    """
    return merge_cells(table, row_index, start_col, row_index, end_col)


def merge_cells_vertical(table, col_index, start_row, end_row):
    """
    Merge cells vertically in a single column.
    
    Args:
        table: The table containing cells to merge
        col_index: Column index (0-based)
        start_row: Starting row index (0-based)
        end_row: Ending row index (0-based, inclusive)
        
    Returns:
        True if successful, False otherwise
    """
    return merge_cells(table, start_row, col_index, end_row, col_index)


def set_cell_alignment(cell, horizontal="left", vertical="top"):
    """
    Set text alignment within a cell.
    
    Args:
        cell: The table cell to format
        horizontal: Horizontal alignment ("left", "center", "right", "justify")
        vertical: Vertical alignment ("top", "center", "bottom")
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Set horizontal alignment for all paragraphs in the cell
        for paragraph in cell.paragraphs:
            if horizontal.lower() == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif horizontal.lower() == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif horizontal.lower() == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:  # default to left
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Set vertical alignment for the cell using XML manipulation
        tc_pr = cell._tc.get_or_add_tcPr()
        
        # Remove existing vertical alignment
        existing_valign = tc_pr.find(qn('w:vAlign'))
        if existing_valign is not None:
            tc_pr.remove(existing_valign)
        
        # Create vertical alignment element
        valign_element = OxmlElement('w:vAlign')
        if vertical.lower() == "center":
            valign_element.set(qn('w:val'), 'center')
        elif vertical.lower() == "bottom":
            valign_element.set(qn('w:val'), 'bottom')
        else:  # default to top
            valign_element.set(qn('w:val'), 'top')
        
        tc_pr.append(valign_element)
        
        return True
        
    except Exception as e:
        print(f"Error setting cell alignment: {e}")
        return False


def set_cell_alignment_by_position(table, row_index, col_index, horizontal="left", vertical="top"):
    """
    Set text alignment for a specific cell by position.
    
    Args:
        table: The table containing the cell
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        horizontal: Horizontal alignment ("left", "center", "right", "justify")
        vertical: Vertical alignment ("top", "center", "bottom")
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if (0 <= row_index < len(table.rows) and 
            0 <= col_index < len(table.rows[row_index].cells)):
            cell = table.rows[row_index].cells[col_index]
            return set_cell_alignment(cell, horizontal, vertical)
        else:
            return False
    except Exception as e:
        print(f"Error setting cell alignment by position: {e}")
        return False


def set_table_alignment(table, horizontal="left", vertical="top"):
    """
    Set text alignment for all cells in a table.
    
    Args:
        table: The table to format
        horizontal: Horizontal alignment ("left", "center", "right", "justify")
        vertical: Vertical alignment ("top", "center", "bottom")
        
    Returns:
        True if successful, False otherwise
    """
    try:
        for row in table.rows:
            for cell in row.cells:
                set_cell_alignment(cell, horizontal, vertical)
        return True
    except Exception as e:
        print(f"Error setting table alignment: {e}")
        return False


def set_column_width(table, col_index, width, width_type="dxa"):
    """
    Set the width of a specific column in a table.
    
    Args:
        table: The table to modify
        col_index: Column index (0-based)
        width: Column width value
        width_type: Width type ("dxa" for points*20, "pct" for percentage*50, "auto")
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Validate column index
        if col_index < 0 or col_index >= len(table.columns):
            return False
        
        # Convert width based on type
        if width_type == "dxa":
            # DXA units (twentieths of a point)
            if isinstance(width, (int, float)):
                width_value = str(int(width * 20))
            else:
                width_value = str(width)
        elif width_type == "pct":
            # Percentage (multiply by 50 for Word format)
            if isinstance(width, (int, float)):
                width_value = str(int(width * 50))
            else:
                width_value = str(width)
        else:
            width_value = str(width)
        
        # Iterate through all rows and set width for cells in the specified column
        for row in table.rows:
            if col_index < len(row.cells):
                cell = row.cells[col_index]
                tc_pr = cell._tc.get_or_add_tcPr()
                
                # Remove existing width
                existing_width = tc_pr.find(qn('w:tcW'))
                if existing_width is not None:
                    tc_pr.remove(existing_width)
                
                # Create new width element
                width_element = OxmlElement('w:tcW')
                width_element.set(qn('w:w'), width_value)
                width_element.set(qn('w:type'), width_type)
                
                tc_pr.append(width_element)
        
        return True
        
    except Exception as e:
        print(f"Error setting column width: {e}")
        return False


def set_column_width_by_position(table, col_index, width, width_type="dxa"):
    """
    Set the width of a specific column by position.
    
    Args:
        table: The table containing the column
        col_index: Column index (0-based)
        width: Column width value
        width_type: Width type ("dxa" for points*20, "pct" for percentage*50, "auto")
        
    Returns:
        True if successful, False otherwise
    """
    return set_column_width(table, col_index, width, width_type)


def set_column_widths(table, widths, width_type="dxa"):
    """
    Set widths for multiple columns in a table.
    
    Args:
        table: The table to modify
        widths: List of width values for each column
        width_type: Width type ("dxa" for points*20, "pct" for percentage*50, "auto")
        
    Returns:
        True if successful, False otherwise
    """
    try:
        for col_index, width in enumerate(widths):
            if col_index >= len(table.columns):
                break
            if not set_column_width(table, col_index, width, width_type):
                return False
        return True
    except Exception as e:
        print(f"Error setting column widths: {e}")
        return False


def set_table_width(table, width, width_type="dxa"):
    """
    Set the overall width of a table.
    
    Args:
        table: The table to modify
        width: Table width value
        width_type: Width type ("dxa" for points*20, "pct" for percentage*50, "auto")
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Convert width based on type
        if width_type == "dxa":
            # DXA units (twentieths of a point)
            if isinstance(width, (int, float)):
                width_value = str(int(width * 20))
            else:
                width_value = str(width)
        elif width_type == "pct":
            # Percentage (multiply by 50 for Word format)
            if isinstance(width, (int, float)):
                width_value = str(int(width * 50))
            else:
                width_value = str(width)
        else:
            width_value = str(width)
        
        # Get table element and properties
        tbl = table._tbl
        
        # Get or create table properties
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        
        # Remove existing table width
        existing_width = tbl_pr.find(qn('w:tblW'))
        if existing_width is not None:
            tbl_pr.remove(existing_width)
        
        # Create new table width element
        width_element = OxmlElement('w:tblW')
        width_element.set(qn('w:w'), width_value)
        width_element.set(qn('w:type'), width_type)
        
        tbl_pr.append(width_element)
        
        return True
        
    except Exception as e:
        print(f"Error setting table width: {e}")
        return False


def auto_fit_table(table):
    """
    Set table to auto-fit columns based on content.
    
    Args:
        table: The table to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Get table element and properties
        tbl = table._tbl
        
        # Get or create table properties
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        
        # Remove existing layout
        existing_layout = tbl_pr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tbl_pr.remove(existing_layout)
        
        # Create auto layout element
        layout_element = OxmlElement('w:tblLayout')
        layout_element.set(qn('w:type'), 'autofit')
        
        tbl_pr.append(layout_element)
        
        # Set all column widths to auto
        for col_index in range(len(table.columns)):
            set_column_width(table, col_index, 0, "auto")
        
        return True
        
    except Exception as e:
        print(f"Error setting auto-fit table: {e}")
        return False


def format_cell_text(cell, text_content=None, bold=None, italic=None, underline=None, 
                    color=None, font_size=None, font_name=None):
    """
    Format text within a table cell.
    
    Args:
        cell: The table cell to format
        text_content: Optional new text content for the cell
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (hex string like "FF0000" or color name)
        font_size: Font size in points
        font_name: Font name/family
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Set text content if provided
        if text_content is not None:
            cell.text = str(text_content)
        
        # Apply formatting to all paragraphs and runs in the cell
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                    
                if font_size is not None:
                    from docx.shared import Pt
                    run.font.size = Pt(font_size)
                    
                if font_name is not None:
                    run.font.name = font_name
                    
                if color is not None:
                    from docx.shared import RGBColor
                    # Define common RGB colors
                    color_map = {
                        'red': RGBColor(255, 0, 0),
                        'blue': RGBColor(0, 0, 255),
                        'green': RGBColor(0, 128, 0),
                        'yellow': RGBColor(255, 255, 0),
                        'black': RGBColor(0, 0, 0),
                        'gray': RGBColor(128, 128, 128),
                        'grey': RGBColor(128, 128, 128),
                        'white': RGBColor(255, 255, 255),
                        'purple': RGBColor(128, 0, 128),
                        'orange': RGBColor(255, 165, 0)
                    }
                    
                    try:
                        if color.lower() in color_map:
                            # Use predefined RGB color
                            run.font.color.rgb = color_map[color.lower()]
                        elif color.startswith('#'):
                            # Hex color string
                            hex_color = color.lstrip('#')
                            if len(hex_color) == 6:
                                r = int(hex_color[0:2], 16)
                                g = int(hex_color[2:4], 16)
                                b = int(hex_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                        else:
                            # Try hex without #
                            if len(color) == 6:
                                r = int(color[0:2], 16)
                                g = int(color[2:4], 16)
                                b = int(color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                    except Exception:
                        # If color parsing fails, default to black
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        return True
        
    except Exception as e:
        print(f"Error formatting cell text: {e}")
        return False


def format_cell_text_by_position(table, row_index, col_index, text_content=None, 
                                 bold=None, italic=None, underline=None, color=None, 
                                 font_size=None, font_name=None):
    """
    Format text in a specific table cell by position.
    
    Args:
        table: The table containing the cell
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        text_content: Optional new text content for the cell
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (hex string or color name)
        font_size: Font size in points
        font_name: Font name/family
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if (0 <= row_index < len(table.rows) and 
            0 <= col_index < len(table.rows[row_index].cells)):
            cell = table.rows[row_index].cells[col_index]
            return format_cell_text(cell, text_content, bold, italic, underline, 
                                   color, font_size, font_name)
        else:
            return False
    except Exception as e:
        print(f"Error formatting cell text by position: {e}")
        return False


def set_cell_padding(cell, top=None, bottom=None, left=None, right=None, unit="dxa"):
    """
    Set padding/margins for a table cell.
    
    Args:
        cell: The table cell to format
        top: Top padding value
        bottom: Bottom padding value
        left: Left padding value
        right: Right padding value
        unit: Unit type ("dxa" for twentieths of a point, "pct" for percentage)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Get or create table cell properties
        tc_pr = cell._tc.get_or_add_tcPr()
        
        # Remove existing margins
        existing_margins = tc_pr.find(qn('w:tcMar'))
        if existing_margins is not None:
            tc_pr.remove(existing_margins)
        
        # Create margins element if any padding is specified
        if any(p is not None for p in [top, bottom, left, right]):
            margins_element = OxmlElement('w:tcMar')
            
            # Add individual margin elements
            margin_sides = {
                'w:top': top,
                'w:bottom': bottom,
                'w:left': left,
                'w:right': right
            }
            
            for side, value in margin_sides.items():
                if value is not None:
                    margin_el = OxmlElement(side)
                    if unit == "dxa":
                        # DXA units (twentieths of a point)
                        margin_el.set(qn('w:w'), str(int(value * 20)))
                        margin_el.set(qn('w:type'), 'dxa')
                    elif unit == "pct":
                        # Percentage
                        margin_el.set(qn('w:w'), str(int(value * 50)))
                        margin_el.set(qn('w:type'), 'pct')
                    else:
                        # Default to DXA
                        margin_el.set(qn('w:w'), str(int(value * 20)))
                        margin_el.set(qn('w:type'), 'dxa')
                    
                    margins_element.append(margin_el)
            
            tc_pr.append(margins_element)
        
        return True
        
    except Exception as e:
        print(f"Error setting cell padding: {e}")
        return False


def set_cell_padding_by_position(table, row_index, col_index, top=None, bottom=None, 
                                left=None, right=None, unit="dxa"):
    """
    Set padding for a specific table cell by position.
    
    Args:
        table: The table containing the cell
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        top: Top padding value
        bottom: Bottom padding value
        left: Left padding value
        right: Right padding value
        unit: Unit type ("dxa" for twentieths of a point, "pct" for percentage)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if (0 <= row_index < len(table.rows) and 
            0 <= col_index < len(table.rows[row_index].cells)):
            cell = table.rows[row_index].cells[col_index]
            return set_cell_padding(cell, top, bottom, left, right, unit)
        else:
            return False
    except Exception as e:
        print(f"Error setting cell padding by position: {e}")
        return False
