"""
Document protection functionality for Word Document Server.
"""
import os
import json
import hashlib
import datetime
from typing import Dict, List, Tuple, Optional, Any


def add_protection_info(doc_path: str, protection_type: str, password_hash: str, 
                        sections: Optional[List[str]] = None, 
                        signature_info: Optional[Dict[str, Any]] = None,
                        raw_password: Optional[str] = None) -> bool:
    """
    Add document protection information to a separate metadata file and encrypt the document.
    
    Args:
        doc_path: Path to the document
        protection_type: Type of protection ('password', 'restricted', 'signature')
        password_hash: Hashed password for security
        sections: List of section names that can be edited (for restricted editing)
        signature_info: Information about digital signature
        raw_password: The actual password for document encryption
        
    Returns:
        True if protection info was successfully added, False otherwise
    """
    # Create metadata filename based on document path
    base_path, _ = os.path.splitext(doc_path)
    metadata_path = f"{base_path}.protection"
    
    # Prepare protection data
    protection_data = {
        "type": protection_type,
        "password_hash": password_hash,
        "applied_date": datetime.datetime.now().isoformat(),
    }
    
    if sections:
        protection_data["editable_sections"] = sections
        
    if signature_info:
        protection_data["signature"] = signature_info
    
    # Write protection info to metadata file
    try:
        with open(metadata_path, 'w') as f:
            json.dump(protection_data, f, indent=2)
        
        # Apply actual document encryption if raw_password is provided
        if protection_type == "password" and raw_password:
            import msoffcrypto
            import tempfile
            import shutil
            
            # Create a temporary file for the encrypted output
            temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(temp_fd)
            
            try:
                # Open the document
                with open(doc_path, 'rb') as f:
                    office_file = msoffcrypto.OfficeFile(f)
                    
                    # Encrypt with password
                    office_file.load_key(password=raw_password)
                    
                    # Write the encrypted file to the temp path
                    with open(temp_path, 'wb') as out_file:
                        office_file.encrypt(out_file)
                
                # Replace original with encrypted version
                shutil.move(temp_path, doc_path)
                
                # Update metadata to note that true encryption was applied
                protection_data["true_encryption"] = True
                with open(metadata_path, 'w') as f:
                    json.dump(protection_data, f, indent=2)
                    
            except Exception as e:
                print(f"Encryption error: {str(e)}")
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                return False
        
        return True
    except Exception as e:
        print(f"Protection error: {str(e)}")
        return False


def verify_document_protection(doc_path: str, password: Optional[str] = None) -> Tuple[bool, str]:
    """
    Verify if a document is protected and if the password is correct.
    
    Args:
        doc_path: Path to the document
        password: Password to verify
    
    Returns:
        Tuple of (is_protected_and_verified, message)
    """
    base_path, _ = os.path.splitext(doc_path)
    metadata_path = f"{base_path}.protection"
    
    # Check if protection metadata exists
    if not os.path.exists(metadata_path):
        return False, "Document is not protected"
    
    try:
        # Read protection data
        with open(metadata_path, 'r') as f:
            protection_data = json.load(f)
        
        # If password is provided, verify it
        if password:
            password_hash = hashlib.sha256(password.encode()).hexdigest()
            if password_hash != protection_data.get("password_hash"):
                return False, "Incorrect password"
        
        # Return protection type
        protection_type = protection_data.get("type", "unknown")
        return True, f"Document is protected with {protection_type} protection"
        
    except Exception as e:
        return False, f"Error verifying protection: {str(e)}"


def is_section_editable(doc_path: str, section_name: str) -> bool:
    """
    Check if a specific section of a document is editable.
    
    Args:
        doc_path: Path to the document
        section_name: Name of the section to check
    
    Returns:
        True if section is editable, False otherwise
    """
    base_path, _ = os.path.splitext(doc_path)
    metadata_path = f"{base_path}.protection"
    
    # Check if protection metadata exists
    if not os.path.exists(metadata_path):
        # If no protection exists, all sections are editable
        return True
    
    try:
        # Read protection data
        with open(metadata_path, 'r') as f:
            protection_data = json.load(f)
        
        # Check protection type
        if protection_data.get("type") != "restricted":
            # If not restricted editing, return based on protection type
            return protection_data.get("type") != "password"
        
        # Check if the section is in the list of editable sections
        editable_sections = protection_data.get("editable_sections", [])
        return section_name in editable_sections
        
    except Exception:
        # In case of error, default to not editable for security
        return False


def create_signature_info(doc, signer_name: str, reason: Optional[str] = None) -> Dict[str, Any]:
    """
    Create signature information for a document.
    
    Args:
        doc: Document object
        signer_name: Name of the person signing the document
        reason: Optional reason for signing
        
    Returns:
        Dictionary containing signature information
    """
    # Create signature info
    signature_info = {
        "signer": signer_name,
        "timestamp": datetime.datetime.now().isoformat(),
    }
    
    if reason:
        signature_info["reason"] = reason
    
    # Generate a simple signature hash based on document content and metadata
    text_content = "\n".join([p.text for p in doc.paragraphs])
    content_hash = hashlib.sha256(text_content.encode()).hexdigest()
    signature_info["content_hash"] = content_hash
    
    return signature_info


def verify_signature(doc_path: str) -> Tuple[bool, str]:
    """
    Verify a document's digital signature.
    
    Args:
        doc_path: Path to the document
        
    Returns:
        Tuple of (is_valid, message)
    """
    from docx import Document
    
    base_path, _ = os.path.splitext(doc_path)
    metadata_path = f"{base_path}.protection"
    
    if not os.path.exists(metadata_path):
        return False, "Document is not signed"
    
    try:
        # Read protection data
        with open(metadata_path, 'r') as f:
            protection_data = json.load(f)
        
        if protection_data.get("type") != "signature":
            return False, f"Document is protected with {protection_data.get('type')} protection, not a signature"
        
        # Get the original content hash
        signature_info = protection_data.get("signature", {})
        original_hash = signature_info.get("content_hash")
        
        if not original_hash:
            return False, "Invalid signature: missing content hash"
        
        # Calculate current content hash
        doc = Document(doc_path)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        current_hash = hashlib.sha256(text_content.encode()).hexdigest()
        
        # Compare hashes
        if current_hash != original_hash:
            return False, f"Document has been modified since it was signed by {signature_info.get('signer')}"
        
        return True, f"Document signature is valid. Signed by {signature_info.get('signer')} on {signature_info.get('timestamp')}"
    
    except Exception as e:
        return False, f"Error verifying signature: {str(e)}"
