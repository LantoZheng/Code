"""
Document creation and manipulation tools for Word Document Server.
"""
import os
import json
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, create_document_copy
from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure, get_document_xml, insert_header_near_text, insert_line_or_paragraph_near_text
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    filename = ensure_docx_extension(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get information about a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    return extract_document_text(filename)


async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)
    
    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"


async def merge_documents(target_filename: str, source_filenames: List[str], add_page_breaks: bool = True) -> str:
    """Merge multiple Word documents into a single document.
    
    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
    """
    from word_document_server.core.tables import copy_table
    
    target_filename = ensure_docx_extension(target_filename)
    
    # Check if target file is writeable
    is_writeable, error_message = check_file_writeable(target_filename)
    if not is_writeable:
        return f"Cannot create target document: {error_message}"
    
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)
    
    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"
    
    try:
        # Create a new document for the merged result
        target_doc = Document()
        
        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc = Document(doc_filename)
            
            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()
            
            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles['Normal']  # Default style
                
                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass
                
                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size
            
            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)
        
        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"


async def get_document_xml_tool(filename: str) -> str:
    """Get the raw XML structure of a Word document."""
    return get_document_xml(filename)
