"""
Comment extraction tools for Word Document Server.

These tools provide high-level interfaces for extracting and analyzing
comments from Word documents through the MCP protocol.
"""
import os
import json
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import ensure_docx_extension
from word_document_server.core.comments import (
    extract_all_comments,
    filter_comments_by_author,
    get_comments_for_paragraph
)


async def get_all_comments(filename: str) -> str:
    """
    Extract all comments from a Word document.
    
    Args:
        filename: Path to the Word document
        
    Returns:
        JSON string containing all comments with metadata
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        comments = extract_all_comments(doc)
        
        # Return results
        return json.dumps({
            'success': True,
            'comments': comments,
            'total_comments': len(comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_by_author(filename: str, author: str) -> str:
    """
    Extract comments from a specific author in a Word document.
    
    Args:
        filename: Path to the Word document
        author: Name of the comment author to filter by
        
    Returns:
        JSON string containing filtered comments
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    if not author or not author.strip():
        return json.dumps({
            'success': False,
            'error': 'Author name cannot be empty'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter by author
        author_comments = filter_comments_by_author(all_comments, author)
        
        # Return results
        return json.dumps({
            'success': True,
            'author': author,
            'comments': author_comments,
            'total_comments': len(author_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_for_paragraph(filename: str, paragraph_index: int) -> str:
    """
    Extract comments for a specific paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        
    Returns:
        JSON string containing comments for the specified paragraph
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return json.dumps({
            'success': False,
            'error': f'Document {filename} does not exist'
        }, indent=2)
    
    if paragraph_index < 0:
        return json.dumps({
            'success': False,
            'error': 'Paragraph index must be non-negative'
        }, indent=2)
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Check if paragraph index is valid
        if paragraph_index >= len(doc.paragraphs):
            return json.dumps({
                'success': False,
                'error': f'Paragraph index {paragraph_index} is out of range. Document has {len(doc.paragraphs)} paragraphs.'
            }, indent=2)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter for the specific paragraph
        from word_document_server.core.comments import get_comments_for_paragraph as core_get_comments_for_paragraph
        para_comments = core_get_comments_for_paragraph(all_comments, paragraph_index)
        
        # Get the paragraph text for context
        paragraph_text = doc.paragraphs[paragraph_index].text
        
        # Return results
        return json.dumps({
            'success': True,
            'paragraph_index': paragraph_index,
            'paragraph_text': paragraph_text,
            'comments': para_comments,
            'total_comments': len(para_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)