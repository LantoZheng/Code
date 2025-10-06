"""
Footnote and endnote tools for Word Document Server.

These tools handle footnote and endnote functionality,
including adding, customizing, and converting between them.

This module combines both standard and robust implementations:
- String-return functions for backward compatibility
- Dict-return robust functions for structured responses
"""
import os
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.core.footnotes import (
    find_footnote_references,
    get_format_symbols,
    customize_footnote_formatting,
    add_footnote_robust,
    delete_footnote_robust,
    validate_document_footnotes,
    add_footnote_at_paragraph_end  # Compatibility function
)


async def add_footnote_to_document(filename: str, paragraph_index: int, footnote_text: str) -> str:
    """Add a footnote to a specific paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to add footnote to (0-based)
        footnote_text: Text content of the footnote
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure paragraph_index is an integer
    try:
        paragraph_index = int(paragraph_index)
    except (ValueError, TypeError):
        return "Invalid parameter: paragraph_index must be an integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        paragraph = doc.paragraphs[paragraph_index]
        
        # In python-docx, we'd use paragraph.add_footnote(), but we'll use a more robust approach
        try:
            footnote = paragraph.add_run()
            footnote.text = ""
            
            # Create the footnote reference
            reference = footnote.add_footnote(footnote_text)
            
            doc.save(filename)
            return f"Footnote added to paragraph {paragraph_index} in {filename}"
        except AttributeError:
            # Fall back to a simpler approach if direct footnote addition fails
            last_run = paragraph.add_run()
            last_run.text = "¹"  # Unicode superscript 1
            last_run.font.superscript = True
            
            # Add a footnote section at the end if it doesn't exist
            found_footnote_section = False
            for p in doc.paragraphs:
                if p.text.startswith("Footnotes:"):
                    found_footnote_section = True
                    break
            
            if not found_footnote_section:
                doc.add_paragraph("\n").add_run()
                doc.add_paragraph("Footnotes:").bold = True
            
            # Add footnote text
            footnote_para = doc.add_paragraph("¹ " + footnote_text)
            footnote_para.style = "Footnote Text" if "Footnote Text" in doc.styles else "Normal"
            
            doc.save(filename)
            return f"Footnote added to paragraph {paragraph_index} in {filename} (simplified approach)"
    except Exception as e:
        return f"Failed to add footnote: {str(e)}"


async def add_endnote_to_document(filename: str, paragraph_index: int, endnote_text: str) -> str:
    """Add an endnote to a specific paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to add endnote to (0-based)
        endnote_text: Text content of the endnote
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure paragraph_index is an integer
    try:
        paragraph_index = int(paragraph_index)
    except (ValueError, TypeError):
        return "Invalid parameter: paragraph_index must be an integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        paragraph = doc.paragraphs[paragraph_index]
        
        # Add endnote reference
        last_run = paragraph.add_run()
        last_run.text = "†"  # Unicode dagger symbol common for endnotes
        last_run.font.superscript = True
        
        # Check if endnotes section exists, if not create it
        endnotes_heading_found = False
        for para in doc.paragraphs:
            if para.text == "Endnotes:" or para.text == "ENDNOTES":
                endnotes_heading_found = True
                break
        
        if not endnotes_heading_found:
            # Add a page break before endnotes section
            doc.add_page_break()
            doc.add_heading("Endnotes:", level=1)
        
        # Add the endnote text
        endnote_para = doc.add_paragraph("† " + endnote_text)
        endnote_para.style = "Endnote Text" if "Endnote Text" in doc.styles else "Normal"
        
        doc.save(filename)
        return f"Endnote added to paragraph {paragraph_index} in {filename}"
    except Exception as e:
        return f"Failed to add endnote: {str(e)}"


async def convert_footnotes_to_endnotes_in_document(filename: str) -> str:
    """Convert all footnotes to endnotes in a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
  
      
        # Find all runs that might be footnote references
        footnote_references = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            for run_idx, run in enumerate(para.runs):
                # Check if this run is likely a footnote reference
                # (superscript number or special character)
                if run.font.superscript and (run.text.isdigit() or run.text in "¹²³⁴⁵⁶⁷⁸⁹"):
                    footnote_references.append({
                        "paragraph_index": para_idx,
                        "run_index": run_idx,
                        "text": run.text
                    })
        
        if not footnote_references:
            return f"No footnote references found in {filename}"
        
        # Create endnotes section
        doc.add_page_break()
        doc.add_heading("Endnotes:", level=1)
        
        # Create a placeholder for endnote content, we'll fill it later
        endnote_content = []
        
        # Find the footnote text at the bottom of the page
       
        found_footnote_section = False
        footnote_text = []
        
        for para in doc.paragraphs:
            if not found_footnote_section and para.text.startswith("Footnotes:"):
                found_footnote_section = True
                continue
            
            if found_footnote_section:
                footnote_text.append(para.text)
        
        # Create endnotes based on footnote references
        for i, ref in enumerate(footnote_references):
            # Add a new endnote
            endnote_para = doc.add_paragraph()
            
            # Try to match with footnote text, or use placeholder
            if i < len(footnote_text):
                endnote_para.text = f"†{i+1} {footnote_text[i]}"
            else:
                endnote_para.text = f"†{i+1} Converted from footnote {ref['text']}"
            
            # Change the footnote reference to an endnote reference
            try:
                paragraph = doc.paragraphs[ref["paragraph_index"]]
                paragraph.runs[ref["run_index"]].text = f"†{i+1}"
            except IndexError:
                # Skip if we can't locate the reference
                pass
        
        # Save the document
        doc.save(filename)
        
        return f"Converted {len(footnote_references)} footnotes to endnotes in {filename}"
    except Exception as e:
        return f"Failed to convert footnotes to endnotes: {str(e)}"


async def add_footnote_after_text(filename: str, search_text: str, footnote_text: str, 
                                 output_filename: Optional[str] = None) -> str:
    """Add a footnote after specific text in a Word document with proper formatting.
    
    This enhanced function ensures proper superscript formatting by managing styles at the XML level.
    
    Args:
        filename: Path to the Word document
        search_text: Text to search for (footnote will be added after this text)
        footnote_text: Content of the footnote
        output_filename: Optional output filename (if None, modifies in place)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Use robust implementation
        success, message, details = add_footnote_robust(
            filename=filename,
            search_text=search_text,
            footnote_text=footnote_text,
            output_filename=output_filename,
            position="after",
            validate_location=True
        )
        return message
    except Exception as e:
        return f"Failed to add footnote: {str(e)}"


async def add_footnote_before_text(filename: str, search_text: str, footnote_text: str, 
                                  output_filename: Optional[str] = None) -> str:
    """Add a footnote before specific text in a Word document with proper formatting.
    
    This enhanced function ensures proper superscript formatting by managing styles at the XML level.
    
    Args:
        filename: Path to the Word document
        search_text: Text to search for (footnote will be added before this text)
        footnote_text: Content of the footnote
        output_filename: Optional output filename (if None, modifies in place)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Use robust implementation
        success, message, details = add_footnote_robust(
            filename=filename,
            search_text=search_text,
            footnote_text=footnote_text,
            output_filename=output_filename,
            position="before",
            validate_location=True
        )
        return message
    except Exception as e:
        return f"Failed to add footnote: {str(e)}"


async def add_footnote_enhanced(filename: str, paragraph_index: int, footnote_text: str,
                               output_filename: Optional[str] = None) -> str:
    """Enhanced version of add_footnote_to_document with proper superscript formatting.
    
    Now uses the robust implementation for better reliability.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to add footnote to (0-based)
        footnote_text: Text content of the footnote
        output_filename: Optional output filename (if None, modifies in place)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure paragraph_index is an integer
    try:
        paragraph_index = int(paragraph_index)
    except (ValueError, TypeError):
        return "Invalid parameter: paragraph_index must be an integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Use robust implementation
        success, message, details = add_footnote_robust(
            filename=filename,
            paragraph_index=paragraph_index,
            footnote_text=footnote_text,
            output_filename=output_filename,
            validate_location=True
        )
        return message
    except Exception as e:
        return f"Failed to add footnote: {str(e)}"


async def customize_footnote_style(filename: str, numbering_format: str = "1, 2, 3", 
                                  start_number: int = 1, font_name: Optional[str] = None,
                                  font_size: Optional[int] = None) -> str:
    """Customize footnote numbering and formatting in a Word document.
    
    Args:
        filename: Path to the Word document
        numbering_format: Format for footnote numbers (e.g., "1, 2, 3", "i, ii, iii", "a, b, c")
        start_number: Number to start footnote numbering from
        font_name: Optional font name for footnotes
        font_size: Optional font size for footnotes (in points)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Create or get footnote style
        footnote_style_name = "Footnote Text"
        footnote_style = None
        
        try:
            footnote_style = doc.styles[footnote_style_name]
        except KeyError:
            # Create the style if it doesn't exist
            footnote_style = doc.styles.add_style(footnote_style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        # Apply formatting to footnote style
        if footnote_style:
            if font_name:
                footnote_style.font.name = font_name
            if font_size:
                footnote_style.font.size = Pt(font_size)
        
        # Find all existing footnote references
        footnote_refs = find_footnote_references(doc)
        
        # Generate format symbols for the specified numbering format
        format_symbols = get_format_symbols(numbering_format, len(footnote_refs) + start_number)
        
        # Apply custom formatting to footnotes
        count = customize_footnote_formatting(doc, footnote_refs, format_symbols, start_number, footnote_style)
        
        # Save the document
        doc.save(filename)
        
        return f"Footnote style and numbering customized in {filename}"
    except Exception as e:
        return f"Failed to customize footnote style: {str(e)}"


async def delete_footnote_from_document(filename: str, footnote_id: Optional[int] = None,
                                       search_text: Optional[str] = None, 
                                       output_filename: Optional[str] = None) -> str:
    """Delete a footnote from a Word document.
    
    You can identify the footnote to delete either by:
    1. footnote_id: The numeric ID of the footnote (1, 2, 3, etc.)
    2. search_text: Text near the footnote reference to find and delete
    
    Args:
        filename: Path to the Word document
        footnote_id: Optional ID of the footnote to delete (1-based)
        search_text: Optional text to search near the footnote reference
        output_filename: Optional output filename (if None, modifies in place)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Use robust implementation with orphan cleanup
        success, message, details = delete_footnote_robust(
            filename=filename,
            footnote_id=footnote_id,
            search_text=search_text,
            output_filename=output_filename,
            clean_orphans=True
        )
        return message
    except Exception as e:
        return f"Failed to delete footnote: {str(e)}"


# ============================================================================
# Robust tool functions with Dict returns for structured responses
# ============================================================================


async def add_footnote_robust_tool(
    filename: str,
    search_text: Optional[str] = None,
    paragraph_index: Optional[int] = None,
    footnote_text: str = "",
    validate_location: bool = True,
    auto_repair: bool = False
) -> Dict[str, Any]:
    """
    Add a footnote with robust validation and error handling.
    
    This is the production-ready version with comprehensive Word compliance.
    
    Args:
        filename: Path to the Word document
        search_text: Text to search for (mutually exclusive with paragraph_index)
        paragraph_index: Index of paragraph (mutually exclusive with search_text)
        footnote_text: Content of the footnote
        validate_location: Whether to validate placement restrictions
        auto_repair: Whether to attempt automatic document repair
    
    Returns:
        Dict with success status, message, and optional details
    """
    filename = ensure_docx_extension(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return {
            "success": False,
            "message": f"Cannot modify document: {error_message}",
            "details": None
        }
    
    # Convert paragraph_index if provided as string
    if paragraph_index is not None:
        try:
            paragraph_index = int(paragraph_index)
        except (ValueError, TypeError):
            return {
                "success": False,
                "message": "Invalid parameter: paragraph_index must be an integer",
                "details": None
            }
    
    # Call robust implementation
    success, message, details = add_footnote_robust(
        filename=filename,
        search_text=search_text,
        paragraph_index=paragraph_index,
        footnote_text=footnote_text,
        validate_location=validate_location,
        auto_repair=auto_repair
    )
    
    return {
        "success": success,
        "message": message,
        "details": details
    }


async def delete_footnote_robust_tool(
    filename: str,
    footnote_id: Optional[int] = None,
    search_text: Optional[str] = None,
    clean_orphans: bool = True
) -> Dict[str, Any]:
    """
    Delete a footnote with comprehensive cleanup.
    
    Args:
        filename: Path to the Word document
        footnote_id: ID of footnote to delete
        search_text: Text near footnote reference
        clean_orphans: Whether to remove orphaned content
    
    Returns:
        Dict with success status, message, and optional details
    """
    filename = ensure_docx_extension(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return {
            "success": False,
            "message": f"Cannot modify document: {error_message}",
            "details": None
        }
    
    # Convert footnote_id if provided as string
    if footnote_id is not None:
        try:
            footnote_id = int(footnote_id)
        except (ValueError, TypeError):
            return {
                "success": False,
                "message": "Invalid parameter: footnote_id must be an integer",
                "details": None
            }
    
    # Call robust implementation
    success, message, details = delete_footnote_robust(
        filename=filename,
        footnote_id=footnote_id,
        search_text=search_text,
        clean_orphans=clean_orphans
    )
    
    return {
        "success": success,
        "message": message,
        "details": details
    }


async def validate_footnotes_tool(filename: str) -> Dict[str, Any]:
    """
    Validate all footnotes in a document.
    
    Provides comprehensive validation report including:
    - ID conflicts
    - Orphaned content
    - Missing styles
    - Invalid locations
    - Coherence issues
    
    Args:
        filename: Path to the Word document
    
    Returns:
        Dict with validation status and detailed report
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return {
            "valid": False,
            "message": f"Document {filename} does not exist",
            "report": {}
        }
    
    # Call validation
    is_valid, message, report = validate_document_footnotes(filename)
    
    return {
        "valid": is_valid,
        "message": message,
        "report": report
    }


# ============================================================================
# Compatibility wrappers for robust tools (maintain backward compatibility)
# ============================================================================

async def add_footnote_to_document_robust(
    filename: str, 
    paragraph_index: int, 
    footnote_text: str
) -> str:
    """
    Robust version of add_footnote_to_document.
    Maintains backward compatibility with existing API.
    """
    result = await add_footnote_robust_tool(
        filename=filename,
        paragraph_index=paragraph_index,
        footnote_text=footnote_text
    )
    return result["message"]


async def add_footnote_after_text_robust(
    filename: str,
    search_text: str,
    footnote_text: str,
    output_filename: Optional[str] = None
) -> str:
    """
    Robust version of add_footnote_after_text.
    Maintains backward compatibility with existing API.
    """
    # Handle output filename by copying first if needed
    working_file = filename
    if output_filename:
        import shutil
        shutil.copy2(filename, output_filename)
        working_file = output_filename
    
    result = await add_footnote_robust_tool(
        filename=working_file,
        search_text=search_text,
        footnote_text=footnote_text
    )
    return result["message"]


async def add_footnote_before_text_robust(
    filename: str,
    search_text: str,
    footnote_text: str,
    output_filename: Optional[str] = None
) -> str:
    """
    Robust version of add_footnote_before_text.
    Note: Current robust implementation defaults to 'after' position.
    """
    # Handle output filename
    working_file = filename
    if output_filename:
        import shutil
        shutil.copy2(filename, output_filename)
        working_file = output_filename
    
    result = await add_footnote_robust_tool(
        filename=working_file,
        search_text=search_text,
        footnote_text=footnote_text
    )
    return result["message"]


async def delete_footnote_from_document_robust(
    filename: str,
    footnote_id: Optional[int] = None,
    search_text: Optional[str] = None,
    output_filename: Optional[str] = None
) -> str:
    """
    Robust version of delete_footnote_from_document.
    Maintains backward compatibility with existing API.
    """
    # Handle output filename
    working_file = filename
    if output_filename:
        import shutil
        shutil.copy2(filename, output_filename)
        working_file = output_filename
    
    result = await delete_footnote_robust_tool(
        filename=working_file,
        footnote_id=footnote_id,
        search_text=search_text
    )
    return result["message"]
