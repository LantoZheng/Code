"""
Formatting tools for Word Document Server.

These tools handle formatting operations for Word documents,
including text formatting, table formatting, and custom styles.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.core.styles import create_style
from word_document_server.core.tables import (
    apply_table_style, set_cell_shading_by_position, apply_alternating_row_shading,
    highlight_header_row, merge_cells, merge_cells_horizontal, merge_cells_vertical,
    set_cell_alignment_by_position, set_table_alignment, set_column_width_by_position,
    set_column_widths, set_table_width as set_table_width_func, auto_fit_table,
    format_cell_text_by_position, set_cell_padding_by_position
)


async def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int, 
                     bold: Optional[bool] = None, italic: Optional[bool] = None, 
                     underline: Optional[bool] = None, color: Optional[str] = None,
                     font_size: Optional[int] = None, font_name: Optional[str] = None) -> str:
    """Format a specific range of text within a paragraph.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        start_pos: Start position within the paragraph text
        end_pos: End position within the paragraph text
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        paragraph_index = int(paragraph_index)
        start_pos = int(start_pos)
        end_pos = int(end_pos)
        if font_size is not None:
            font_size = int(font_size)
    except (ValueError, TypeError):
        return "Invalid parameter: paragraph_index, start_pos, end_pos, and font_size must be integers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text
        
        # Validate text positions
        if start_pos < 0 or end_pos > len(text) or start_pos >= end_pos:
            return f"Invalid text positions. Paragraph has {len(text)} characters."
        
        # Get the text to format
        target_text = text[start_pos:end_pos]
        
        # Clear existing runs and create three runs: before, target, after
        for run in paragraph.runs:
            run.clear()
        
        # Add text before target
        if start_pos > 0:
            run_before = paragraph.add_run(text[:start_pos])
        
        # Add target text with formatting
        run_target = paragraph.add_run(target_text)
        if bold is not None:
            run_target.bold = bold
        if italic is not None:
            run_target.italic = italic
        if underline is not None:
            run_target.underline = underline
        if color:
            # Define common RGB colors
            color_map = {
                'red': RGBColor(255, 0, 0),
                'blue': RGBColor(0, 0, 255),
                'green': RGBColor(0, 128, 0),
                'yellow': RGBColor(255, 255, 0),
                'black': RGBColor(0, 0, 0),
                'gray': RGBColor(128, 128, 128),
                'white': RGBColor(255, 255, 255),
                'purple': RGBColor(128, 0, 128),
                'orange': RGBColor(255, 165, 0)
            }
            
            try:
                if color.lower() in color_map:
                    # Use predefined RGB color
                    run_target.font.color.rgb = color_map[color.lower()]
                else:
                    # Try to set color by name
                    run_target.font.color.rgb = RGBColor.from_string(color)
            except Exception as e:
                # If all else fails, default to black
                run_target.font.color.rgb = RGBColor(0, 0, 0)
        if font_size:
            run_target.font.size = Pt(font_size)
        if font_name:
            run_target.font.name = font_name
        
        # Add text after target
        if end_pos < len(text):
            run_after = paragraph.add_run(text[end_pos:])
        
        doc.save(filename)
        return f"Text '{target_text}' formatted successfully in paragraph {paragraph_index}."
    except Exception as e:
        return f"Failed to format text: {str(e)}"


async def create_custom_style(filename: str, style_name: str, 
                             bold: Optional[bool] = None, italic: Optional[bool] = None,
                             font_size: Optional[int] = None, font_name: Optional[str] = None,
                             color: Optional[str] = None, base_style: Optional[str] = None) -> str:
    """Create a custom style in the document.
    
    Args:
        filename: Path to the Word document
        style_name: Name for the new style
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        font_size: Font size in points
        font_name: Font name/family
        color: Text color (e.g., 'red', 'blue')
        base_style: Optional existing style to base this on
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Build font properties dictionary
        font_properties = {}
        if bold is not None:
            font_properties['bold'] = bold
        if italic is not None:
            font_properties['italic'] = italic
        if font_size is not None:
            font_properties['size'] = font_size
        if font_name is not None:
            font_properties['name'] = font_name
        if color is not None:
            font_properties['color'] = color
        
        # Create the style
        new_style = create_style(
            doc, 
            style_name, 
            WD_STYLE_TYPE.PARAGRAPH, 
            base_style=base_style,
            font_properties=font_properties
        )
        
        doc.save(filename)
        return f"Style '{style_name}' created successfully."
    except Exception as e:
        return f"Failed to create style: {str(e)}"


async def format_table(filename: str, table_index: int, 
                      has_header_row: Optional[bool] = None,
                      border_style: Optional[str] = None,
                      shading: Optional[List[List[str]]] = None) -> str:
    """Format a table with borders, shading, and structure.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick')
        shading: 2D list of cell background colors (by row and column)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply formatting
        success = apply_table_style(table, has_header_row or False, border_style, shading)
        
        if success:
            doc.save(filename)
            return f"Table at index {table_index} formatted successfully."
        else:
            return f"Failed to format table at index {table_index}."
    except Exception as e:
        return f"Failed to format table: {str(e)}"


async def set_table_cell_shading(filename: str, table_index: int, row_index: int, 
                                col_index: int, fill_color: str, pattern: str = "clear") -> str:
    """Apply shading/filling to a specific table cell.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index of the cell (0-based)
        col_index: Column index of the cell (0-based)
        fill_color: Background color (hex string like "FF0000" or "red")
        pattern: Shading pattern ("clear", "solid", "pct10", "pct20", etc.)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index, row_index, and col_index must be integers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Validate row and column indices
        if row_index < 0 or row_index >= len(table.rows):
            return f"Invalid row index. Table has {len(table.rows)} rows (0-{len(table.rows)-1})."
        
        if col_index < 0 or col_index >= len(table.rows[row_index].cells):
            return f"Invalid column index. Row has {len(table.rows[row_index].cells)} cells (0-{len(table.rows[row_index].cells)-1})."
        
        # Apply cell shading
        success = set_cell_shading_by_position(table, row_index, col_index, fill_color, pattern)
        
        if success:
            doc.save(filename)
            return f"Cell shading applied successfully to table {table_index}, row {row_index}, column {col_index}."
        else:
            return f"Failed to apply cell shading."
    except Exception as e:
        return f"Failed to apply cell shading: {str(e)}"


async def apply_table_alternating_rows(filename: str, table_index: int, 
                                     color1: str = "FFFFFF", color2: str = "F2F2F2") -> str:
    """Apply alternating row colors to a table for better readability.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        color1: Color for odd rows (hex string, default white)
        color2: Color for even rows (hex string, default light gray)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index must be an integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply alternating row shading
        success = apply_alternating_row_shading(table, color1, color2)
        
        if success:
            doc.save(filename)
            return f"Alternating row shading applied successfully to table {table_index}."
        else:
            return f"Failed to apply alternating row shading."
    except Exception as e:
        return f"Failed to apply alternating row shading: {str(e)}"


async def highlight_table_header(filename: str, table_index: int, 
                               header_color: str = "4472C4", text_color: str = "FFFFFF") -> str:
    """Apply special highlighting to table header row.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        header_color: Background color for header (hex string, default blue)
        text_color: Text color for header (hex string, default white)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index must be an integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply header highlighting
        success = highlight_header_row(table, header_color, text_color)
        
        if success:
            doc.save(filename)
            return f"Header highlighting applied successfully to table {table_index}."
        else:
            return f"Failed to apply header highlighting."
    except Exception as e:
        return f"Failed to apply header highlighting: {str(e)}"


async def merge_table_cells(filename: str, table_index: int, start_row: int, start_col: int, 
                          end_row: int, end_col: int) -> str:
    """Merge cells in a rectangular area of a table.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        start_row: Starting row index (0-based)
        start_col: Starting column index (0-based)
        end_row: Ending row index (0-based, inclusive)
        end_col: Ending column index (0-based, inclusive)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        start_row = int(start_row)
        start_col = int(start_col)
        end_row = int(end_row)
        end_col = int(end_col)
    except (ValueError, TypeError):
        return "Invalid parameter: all indices must be integers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Validate merge parameters
        if start_row > end_row or start_col > end_col:
            return "Invalid merge range: start indices must be <= end indices"
        
        if start_row == end_row and start_col == end_col:
            return "Invalid merge range: cannot merge a single cell with itself"
        
        # Apply cell merge
        success = merge_cells(table, start_row, start_col, end_row, end_col)
        
        if success:
            doc.save(filename)
            return f"Cells merged successfully in table {table_index} from ({start_row},{start_col}) to ({end_row},{end_col})."
        else:
            return f"Failed to merge cells. Check that indices are valid."
    except Exception as e:
        return f"Failed to merge cells: {str(e)}"


async def merge_table_cells_horizontal(filename: str, table_index: int, row_index: int, 
                                     start_col: int, end_col: int) -> str:
    """Merge cells horizontally in a single row.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index (0-based)
        start_col: Starting column index (0-based)
        end_col: Ending column index (0-based, inclusive)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        start_col = int(start_col)
        end_col = int(end_col)
    except (ValueError, TypeError):
        return "Invalid parameter: all indices must be integers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply horizontal cell merge
        success = merge_cells_horizontal(table, row_index, start_col, end_col)
        
        if success:
            doc.save(filename)
            return f"Cells merged horizontally in table {table_index}, row {row_index}, columns {start_col}-{end_col}."
        else:
            return f"Failed to merge cells horizontally. Check that indices are valid."
    except Exception as e:
        return f"Failed to merge cells horizontally: {str(e)}"


async def merge_table_cells_vertical(filename: str, table_index: int, col_index: int, 
                                   start_row: int, end_row: int) -> str:
    """Merge cells vertically in a single column.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        col_index: Column index (0-based)
        start_row: Starting row index (0-based)
        end_row: Ending row index (0-based, inclusive)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        col_index = int(col_index)
        start_row = int(start_row)
        end_row = int(end_row)
    except (ValueError, TypeError):
        return "Invalid parameter: all indices must be integers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply vertical cell merge
        success = merge_cells_vertical(table, col_index, start_row, end_row)
        
        if success:
            doc.save(filename)
            return f"Cells merged vertically in table {table_index}, column {col_index}, rows {start_row}-{end_row}."
        else:
            return f"Failed to merge cells vertically. Check that indices are valid."
    except Exception as e:
        return f"Failed to merge cells vertically: {str(e)}"


async def set_table_cell_alignment(filename: str, table_index: int, row_index: int, col_index: int,
                                 horizontal: str = "left", vertical: str = "top") -> str:
    """Set text alignment for a specific table cell.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        horizontal: Horizontal alignment ("left", "center", "right", "justify")
        vertical: Vertical alignment ("top", "center", "bottom")
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index, row_index, and col_index must be integers"
    
    # Validate alignment parameters
    valid_horizontal = ["left", "center", "right", "justify"]
    valid_vertical = ["top", "center", "bottom"]
    
    if horizontal.lower() not in valid_horizontal:
        return f"Invalid horizontal alignment. Valid options: {', '.join(valid_horizontal)}"
    
    if vertical.lower() not in valid_vertical:
        return f"Invalid vertical alignment. Valid options: {', '.join(valid_vertical)}"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply cell alignment
        success = set_cell_alignment_by_position(table, row_index, col_index, horizontal, vertical)
        
        if success:
            doc.save(filename)
            return f"Cell alignment set successfully for table {table_index}, cell ({row_index},{col_index}) to {horizontal}/{vertical}."
        else:
            return f"Failed to set cell alignment. Check that indices are valid."
    except Exception as e:
        return f"Failed to set cell alignment: {str(e)}"


async def set_table_alignment_all(filename: str, table_index: int, 
                                horizontal: str = "left", vertical: str = "top") -> str:
    """Set text alignment for all cells in a table.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        horizontal: Horizontal alignment ("left", "center", "right", "justify")
        vertical: Vertical alignment ("top", "center", "bottom")
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index must be an integer"
    
    # Validate alignment parameters
    valid_horizontal = ["left", "center", "right", "justify"]
    valid_vertical = ["top", "center", "bottom"]
    
    if horizontal.lower() not in valid_horizontal:
        return f"Invalid horizontal alignment. Valid options: {', '.join(valid_horizontal)}"
    
    if vertical.lower() not in valid_vertical:
        return f"Invalid vertical alignment. Valid options: {', '.join(valid_vertical)}"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply table alignment
        success = set_table_alignment(table, horizontal, vertical)
        
        if success:
            doc.save(filename)
            return f"Table alignment set successfully for table {table_index} to {horizontal}/{vertical} for all cells."
        else:
            return f"Failed to set table alignment."
    except Exception as e:
        return f"Failed to set table alignment: {str(e)}"


async def set_table_column_width(filename: str, table_index: int, col_index: int, 
                                width: float, width_type: str = "points") -> str:
    """Set the width of a specific table column.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        col_index: Column index (0-based)
        width: Column width value
        width_type: Width type ("points", "inches", "cm", "percent", "auto")
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        col_index = int(col_index)
        if width_type != "auto":
            width = float(width)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index and col_index must be integers, width must be a number"
    
    # Validate width type
    valid_width_types = ["points", "inches", "cm", "percent", "auto"]
    if width_type.lower() not in valid_width_types:
        return f"Invalid width type. Valid options: {', '.join(valid_width_types)}"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Validate column index
        if col_index < 0 or col_index >= len(table.columns):
            return f"Invalid column index. Table has {len(table.columns)} columns (0-{len(table.columns)-1})."
        
        # Convert width and type for Word format
        if width_type.lower() == "points":
            # Points to DXA (twentieths of a point)
            word_width = width
            word_type = "dxa"
        elif width_type.lower() == "inches":
            # Inches to points, then to DXA
            word_width = width * 72  # 72 points per inch
            word_type = "dxa"
        elif width_type.lower() == "cm":
            # CM to points, then to DXA
            word_width = width * 28.35  # ~28.35 points per cm
            word_type = "dxa"
        elif width_type.lower() == "percent":
            # Percentage (Word uses 50x the percentage value)
            word_width = width
            word_type = "pct"
        else:  # auto
            word_width = 0
            word_type = "auto"
        
        # Apply column width
        success = set_column_width_by_position(table, col_index, word_width, word_type)
        
        if success:
            doc.save(filename)
            return f"Column width set successfully for table {table_index}, column {col_index} to {width} {width_type}."
        else:
            return f"Failed to set column width. Check that indices are valid."
    except Exception as e:
        return f"Failed to set column width: {str(e)}"


async def set_table_column_widths(filename: str, table_index: int, widths: list, 
                                 width_type: str = "points") -> str:
    """Set the widths of multiple table columns.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        widths: List of width values for each column
        width_type: Width type ("points", "inches", "cm", "percent", "auto")
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        if width_type != "auto":
            widths = [float(w) for w in widths]
    except (ValueError, TypeError):
        return "Invalid parameter: table_index must be an integer, widths must be a list of numbers"
    
    # Validate width type
    valid_width_types = ["points", "inches", "cm", "percent", "auto"]
    if width_type.lower() not in valid_width_types:
        return f"Invalid width type. Valid options: {', '.join(valid_width_types)}"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Convert widths and type for Word format
        word_widths = []
        for width in widths:
            if width_type.lower() == "points":
                word_widths.append(width)
            elif width_type.lower() == "inches":
                word_widths.append(width * 72)  # 72 points per inch
            elif width_type.lower() == "cm":
                word_widths.append(width * 28.35)  # ~28.35 points per cm
            elif width_type.lower() == "percent":
                word_widths.append(width)
            else:  # auto
                word_widths.append(0)
        
        # Determine Word type
        if width_type.lower() == "percent":
            word_type = "pct"
        elif width_type.lower() == "auto":
            word_type = "auto"
        else:
            word_type = "dxa"
        
        # Apply column widths
        success = set_column_widths(table, word_widths, word_type)
        
        if success:
            doc.save(filename)
            return f"Column widths set successfully for table {table_index} with {len(widths)} columns in {width_type}."
        else:
            return f"Failed to set column widths."
    except Exception as e:
        return f"Failed to set column widths: {str(e)}"


async def set_table_width(filename: str, table_index: int, width: float, 
                         width_type: str = "points") -> str:
    """Set the overall width of a table.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        width: Table width value
        width_type: Width type ("points", "inches", "cm", "percent", "auto")
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        if width_type != "auto":
            width = float(width)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index must be an integer, width must be a number"
    
    # Validate width type
    valid_width_types = ["points", "inches", "cm", "percent", "auto"]
    if width_type.lower() not in valid_width_types:
        return f"Invalid width type. Valid options: {', '.join(valid_width_types)}"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Convert width and type for Word format
        if width_type.lower() == "points":
            word_width = width
            word_type = "dxa"
        elif width_type.lower() == "inches":
            word_width = width * 72  # 72 points per inch
            word_type = "dxa"
        elif width_type.lower() == "cm":
            word_width = width * 28.35  # ~28.35 points per cm
            word_type = "dxa"
        elif width_type.lower() == "percent":
            word_width = width
            word_type = "pct"
        else:  # auto
            word_width = 0
            word_type = "auto"
        
        # Apply table width
        success = set_table_width_func(table, word_width, word_type)
        
        if success:
            doc.save(filename)
            return f"Table width set successfully for table {table_index} to {width} {width_type}."
        else:
            return f"Failed to set table width."
    except Exception as e:
        return f"Failed to set table width: {str(e)}"


async def auto_fit_table_columns(filename: str, table_index: int) -> str:
    """Set table columns to auto-fit based on content.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index must be an integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply auto-fit
        success = auto_fit_table(table)
        
        if success:
            doc.save(filename)
            return f"Table {table_index} set to auto-fit columns based on content."
        else:
            return f"Failed to set table auto-fit."
    except Exception as e:
        return f"Failed to set table auto-fit: {str(e)}"


async def format_table_cell_text(filename: str, table_index: int, row_index: int, col_index: int,
                                 text_content: Optional[str] = None, bold: Optional[bool] = None, italic: Optional[bool] = None,
                                 underline: Optional[bool] = None, color: Optional[str] = None, font_size: Optional[int] = None,
                                 font_name: Optional[str] = None) -> str:
    """Format text within a specific table cell.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        text_content: Optional new text content for the cell
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (hex string like "FF0000" or color name like "red")
        font_size: Font size in points
        font_name: Font name/family
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
        if font_size is not None:
            font_size = int(font_size)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index, row_index, col_index must be integers, font_size must be integer"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Validate row and column indices
        if row_index < 0 or row_index >= len(table.rows):
            return f"Invalid row index. Table has {len(table.rows)} rows (0-{len(table.rows)-1})."
        
        if col_index < 0 or col_index >= len(table.rows[row_index].cells):
            return f"Invalid column index. Row has {len(table.rows[row_index].cells)} cells (0-{len(table.rows[row_index].cells)-1})."
        
        # Apply cell text formatting
        success = format_cell_text_by_position(table, row_index, col_index, text_content, 
                                              bold, italic, underline, color, font_size, font_name)
        
        if success:
            doc.save(filename)
            format_desc = []
            if text_content is not None:
                format_desc.append(f"content='{text_content[:30]}{'...' if len(text_content) > 30 else ''}'")
            if bold is not None:
                format_desc.append(f"bold={bold}")
            if italic is not None:
                format_desc.append(f"italic={italic}")
            if underline is not None:
                format_desc.append(f"underline={underline}")
            if color is not None:
                format_desc.append(f"color={color}")
            if font_size is not None:
                format_desc.append(f"size={font_size}pt")
            if font_name is not None:
                format_desc.append(f"font={font_name}")
            
            format_str = ", ".join(format_desc) if format_desc else "no changes"
            return f"Cell text formatted successfully in table {table_index}, cell ({row_index},{col_index}): {format_str}."
        else:
            return f"Failed to format cell text. Check that indices are valid."
    except Exception as e:
        return f"Failed to format cell text: {str(e)}"


async def set_table_cell_padding(filename: str, table_index: int, row_index: int, col_index: int,
                                 top: Optional[float] = None, bottom: Optional[float] = None, left: Optional[float] = None, 
                                 right: Optional[float] = None, unit: str = "points") -> str:
    """Set padding/margins for a specific table cell.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        top: Top padding in specified units
        bottom: Bottom padding in specified units
        left: Left padding in specified units
        right: Right padding in specified units
        unit: Unit type ("points" or "percent")
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        row_index = int(row_index)
        col_index = int(col_index)
        if top is not None:
            top = float(top)
        if bottom is not None:
            bottom = float(bottom)
        if left is not None:
            left = float(left)
        if right is not None:
            right = float(right)
    except (ValueError, TypeError):
        return "Invalid parameter: indices must be integers, padding values must be numbers"
    
    # Validate unit
    valid_units = ["points", "percent"]
    if unit.lower() not in valid_units:
        return f"Invalid unit. Valid options: {', '.join(valid_units)}"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Validate row and column indices
        if row_index < 0 or row_index >= len(table.rows):
            return f"Invalid row index. Table has {len(table.rows)} rows (0-{len(table.rows)-1})."
        
        if col_index < 0 or col_index >= len(table.rows[row_index].cells):
            return f"Invalid column index. Row has {len(table.rows[row_index].cells)} cells (0-{len(table.rows[row_index].cells)-1})."
        
        # Convert unit for Word format
        word_unit = "dxa" if unit.lower() == "points" else "pct"
        
        # Apply cell padding
        success = set_cell_padding_by_position(table, row_index, col_index, top, bottom, 
                                              left, right, word_unit)
        
        if success:
            doc.save(filename)
            padding_desc = []
            if top is not None:
                padding_desc.append(f"top={top}")
            if bottom is not None:
                padding_desc.append(f"bottom={bottom}")
            if left is not None:
                padding_desc.append(f"left={left}")
            if right is not None:
                padding_desc.append(f"right={right}")
            
            padding_str = ", ".join(padding_desc) if padding_desc else "no padding"
            return f"Cell padding set successfully for table {table_index}, cell ({row_index},{col_index}): {padding_str} {unit}."
        else:
            return f"Failed to set cell padding. Check that indices are valid."
    except Exception as e:
        return f"Failed to set cell padding: {str(e)}"
